#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档处理服务：负责文档上传、分片、向量化等功能
支持BGE-M3模型进行向量化处理
新增特性：
1. 支持Excel(.xlsx/.xls)文件处理
2. 适配超大文件分片（5000分片上限+批次插入）
3. 自动清洗文本中的图片/视频/Base64/HTML多媒体标签
4. 全量分片长度监控（移除固定索引依赖）
"""

import os
import re
import uuid
import hashlib
import logging
from datetime import datetime
from pymilvus import connections, Collection, utility, DataType
import PyPDF2
from docx import Document
import torch
from transformers import AutoTokenizer, AutoModel
import numpy as np

# 新增Excel处理依赖
try:
    import openpyxl  # 处理.xlsx
    from openpyxl.utils.exceptions import InvalidFileException
except ImportError:
    logging.warning("⚠️ openpyxl未安装，将无法处理.xlsx文件，请执行: pip install openpyxl")
    openpyxl = None

try:
    import xlrd  # 处理.xls
    from xlrd.biffh import XLRDError
except ImportError:
    logging.warning("⚠️ xlrd未安装，将无法处理.xls文件，请执行: pip install xlrd==1.2.0")
    xlrd = None

logger = logging.getLogger(__name__)

# 全局常量定义 - 统一长度限制
MILVUS_CHUNK_CONTENT_MAX_LENGTH = 2000  # Milvus字段最大长度
SAFE_TRUNCATE_LENGTH = 1950  # 安全截断长度，留出50字符余量
WARNING_THRESHOLD = 1900  # 长度告警阈值，接近阈值时重点监控

# 分片监控配置
MONITOR_CONFIG = {
    'enable_extended_log': True,  # 启用扩展日志
    'log_top_largest_chunks': 10,  # 记录长度最大的N个分片
    'log_abnormal_chunks': True,  # 记录长度异常的分片
    'abnormal_threshold': WARNING_THRESHOLD  # 异常长度阈值
}


class DocumentProcessor:
    def __init__(self):
        self.upload_dir = "uploads/uploads_original"
        os.makedirs(self.upload_dir, exist_ok=True)

        # 加载BGE-M3模型
        self.model_path = r"D:\workspace_python\models\bge-m3"
        self.tokenizer = None
        self.model = None
        self.load_model()

    def load_model(self):
        """加载BGE-M3模型"""
        try:
            if os.path.exists(self.model_path):
                self.tokenizer = AutoTokenizer.from_pretrained(self.model_path)
                self.model = AutoModel.from_pretrained(self.model_path)
                logger.info("✅ BGE-M3模型加载成功")
            else:
                logger.warning("⚠️ BGE-M3模型路径不存在，使用默认模型")
                self.tokenizer = AutoTokenizer.from_pretrained("BAAI/bge-m3")
                self.model = AutoModel.from_pretrained("BAAI/bge-m3")
        except Exception as e:
            logger.error(f"❌ 模型加载失败: {str(e)}")

    def get_file_md5(self, file_path):
        """计算文件MD5"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

    def extract_text_from_pdf(self, file_path):
        """从PDF文件中提取文本"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"PDF文本提取失败: {str(e)}")
            return ""

    def extract_text_from_docx(self, file_path):
        """从Word文档中提取文本"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # 提取表格内容（避免表格中的图片标记残留）
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Word文档文本提取失败: {str(e)}")
            return ""

    def extract_text_from_txt(self, file_path):
        """从TXT文件中提取文本"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"TXT文件读取失败: {str(e)}")
            return ""

    def extract_text_from_xlsx(self, file_path):
        """从.xlsx文件中提取文本"""
        if openpyxl is None:
            logger.error("❌ openpyxl未安装，无法处理.xlsx文件")
            return ""

        try:
            text = ""
            # 打开Excel文件（仅读取数据，忽略公式）
            workbook = openpyxl.load_workbook(file_path, data_only=True)

            # 遍历所有工作表
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"=== 工作表: {sheet_name} ===\n"

                # 遍历所有行和列（跳过图片/图表所在单元格）
                for row in sheet.iter_rows(values_only=True):
                    # 过滤空值和图片标记，拼接行内容
                    row_data = []
                    for cell in row:
                        if cell is not None:
                            cell_str = str(cell).strip()
                            # 过滤图片相关标记
                            if not cell_str.startswith("Picture") and not cell_str.startswith("Image"):
                                row_data.append(cell_str)
                    if row_data:
                        text += " | ".join(row_data) + "\n"

                text += "\n"  # 工作表之间换行分隔

            workbook.close()
            return text.strip()
        except InvalidFileException:
            logger.error(f"❌ 无效的.xlsx文件: {file_path}")
            return ""
        except Exception as e:
            logger.error(f"xlsx文本提取失败: {str(e)}")
            return ""

    def extract_text_from_xls(self, file_path):
        """从.xls文件中提取文本"""
        if xlrd is None:
            logger.error("❌ xlrd未安装，无法处理.xls文件")
            return ""

        try:
            text = ""
            # 打开Excel文件
            workbook = xlrd.open_workbook(file_path)

            # 遍历所有工作表
            for sheet_name in workbook.sheet_names():
                sheet = workbook.sheet_by_name(sheet_name)
                text += f"=== 工作表: {sheet_name} ===\n"

                # 遍历所有行和列（跳过图片单元格）
                for row_idx in range(sheet.nrows):
                    row = sheet.row_values(row_idx)
                    # 过滤空值和图片标记，拼接行内容
                    row_data = []
                    for cell in row:
                        if cell is not None:
                            cell_str = str(cell).strip()
                            # 过滤图片相关标记
                            if not cell_str.startswith("Picture") and not cell_str.startswith("Image"):
                                row_data.append(cell_str)
                    if row_data:
                        text += " | ".join(row_data) + "\n"

                text += "\n"  # 工作表之间换行分隔

            return text.strip()
        except XLRDError:
            logger.error(f"❌ 无效的.xls文件: {file_path}")
            return ""
        except Exception as e:
            logger.error(f"xls文本提取失败: {str(e)}")
            return ""

    def clean_media_content(self, text):
        """
        核心功能：清洗文本中的所有图片/视频/多媒体相关内容
        移除：HTML标签、Base64图片、图片/视频标记、二进制媒体内容
        """
        if not text or not isinstance(text, str):
            return ""

        # 1. 移除HTML图片/视频/音频/iframe标签（含嵌套内容）
        text = re.sub(r'<img\b[^>]*?>', '', text, flags=re.IGNORECASE | re.DOTALL)
        text = re.sub(r'<video\b[^>]*?>.*?</video>', '', text, flags=re.IGNORECASE | re.DOTALL)
        text = re.sub(r'<audio\b[^>]*?>.*?</audio>', '', text, flags=re.IGNORECASE | re.DOTALL)
        text = re.sub(r'<iframe\b[^>]*?>.*?</iframe>', '', text, flags=re.IGNORECASE | re.DOTALL)

        # 2. 移除Base64编码的图片内容（PDF/Word提取时常见）
        text = re.sub(r'data:image/(png|jpeg|jpg|gif|bmp|webp);base64,[a-zA-Z0-9+/=]+', '', text, flags=re.IGNORECASE)

        # 3. 移除常见的图片/视频描述标记
        text = re.sub(r'【.*?图片.*?】|【.*?视频.*?】', '', text)  # 中文括号
        text = re.sub(r'\[.*?图片.*?\]|\[.*?视频.*?\]', '', text)  # 英文括号
        text = re.sub(r'Picture\s*\d+|Image\s*\d+|Video\s*\d+', '', text)  # 英文标记

        # 4. 移除二进制乱码（图片/视频二进制内容转文本后的残留）
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\xFF]', '', text)

        # 5. 清理多余空行和空格（清洗后产生的冗余）
        text = re.sub(r'\n\s*\n', '\n', text)
        text = re.sub(r' +', ' ', text)

        return text.strip()

    def safe_truncate_text(self, text, max_length=None):
        """
        安全截断文本，确保长度严格小于等于指定最大值
        :param text: 要截断的文本
        :param max_length: 最大长度，默认使用全局安全截断长度
        :return: 截断后的文本
        """
        if max_length is None:
            max_length = SAFE_TRUNCATE_LENGTH

        if not isinstance(text, str):
            text = str(text) if text else ""

        # 严格截断到指定长度
        if len(text) > max_length:
            truncated_text = text[:max_length]
            logger.debug(f"文本截断: {len(text)} -> {len(truncated_text)} 字符")
            return truncated_text

        return text

    def chunk_text(self, text, chunk_size=500, overlap=50):
        """将文本分片，确保分片内容不超过字段最大长度限制"""
        if not text or len(text.strip()) == 0:
            return []

        chunks = []
        start = 0
        text_length = len(text)

        # 确保overlap不超过chunk_size的一半，避免无限循环
        overlap = min(overlap, chunk_size // 2)

        # 使用全局常量作为最大分片长度
        max_chunk_length = SAFE_TRUNCATE_LENGTH

        while start < text_length:
            # 确保分片长度不超过最大限制
            end = min(start + min(chunk_size, max_chunk_length), text_length)

            # 记录分片范围信息
            logger.debug(f"分片范围: [{start}:{end}], 长度: {end - start}")

            # 如果剩余文本不足chunk_size，直接取剩余部分
            if end - start < chunk_size // 2 and len(chunks) > 0:
                # 检查合并后是否超过最大长度
                merged_length = len(chunks[-1]) + (end - start)
                logger.debug(f"合并检查: 前一个分片长度={len(chunks[-1])}, 合并后长度={merged_length}")
                if merged_length <= max_chunk_length:
                    # 如果合并后不超过限制，合并到最后一个分片
                    chunks[-1] += text[start:end]
                    # 合并后再次检查并截断
                    chunks[-1] = self.safe_truncate_text(chunks[-1], max_chunk_length)
                    logger.debug(f"合并到前一个分片，新长度={len(chunks[-1])}")
                else:
                    # 如果合并后超过限制，创建新的分片
                    chunk = self.safe_truncate_text(text[start:end], max_chunk_length)
                    chunks.append(chunk)
                    logger.debug(f"创建新分片，长度={len(chunk)}")
                break

            chunk = text[start:end]
            # 强制安全截断
            chunk = self.safe_truncate_text(chunk, max_chunk_length)

            chunks.append(chunk)

            # 验证当前分片长度（双重保险）
            if len(chunk) > MILVUS_CHUNK_CONTENT_MAX_LENGTH:
                logger.error(f"严重错误: 分片长度 {len(chunk)} 超过Milvus最大限制 {MILVUS_CHUNK_CONTENT_MAX_LENGTH}")
                chunk = self.safe_truncate_text(chunk, MILVUS_CHUNK_CONTENT_MAX_LENGTH)
                chunks[-1] = chunk

            # 记录当前分片索引和长度
            logger.debug(f"分片 {len(chunks)} 长度: {len(chunk)}")

            # 移动起始位置，考虑重叠
            start = end - overlap

            # 防止无限循环
            if start < 0:
                start = 0

            # 如果已经处理完所有文本，退出循环
            if start >= text_length:
                break

            # 防止重复处理相同位置（针对超大文件的保护）
            if len(chunks) > 5000:  # 提升上限，适配更大文件
                logger.warning(f"分片数量超过5000({len(chunks)}), 强制退出分片过程")
                break

        # 分片完成后输出统计信息
        if chunks:
            chunk_lengths = [len(c) for c in chunks]
            logger.info(f"""
文本分片完成统计:
  - 总分片数: {len(chunks)}
  - 最大分片长度: {max(chunk_lengths)}
  - 最小分片长度: {min(chunk_lengths)}
  - 平均分片长度: {int(np.mean(chunk_lengths))}
  - 超过告警阈值({WARNING_THRESHOLD})的分片数: {sum(1 for l in chunk_lengths if l > WARNING_THRESHOLD)}
            """.strip())
        else:
            logger.info("文本分片完成，共 0 个分片")

        return chunks

    def analyze_chunk_lengths(self, chunks):
        """
        分析分片长度分布，返回统计信息和异常分片列表
        适用于处理超大文件时的分片监控
        """
        if not chunks:
            return {
                'total': 0,
                'max_length': 0,
                'min_length': 0,
                'avg_length': 0,
                'abnormal_chunks': [],
                'top_largest_chunks': []
            }

        # 计算基本统计信息
        chunk_lengths = [(i, len(c.get('chunk_content', ''))) for i, c in enumerate(chunks)]
        total_chunks = len(chunk_lengths)
        lengths = [l for _, l in chunk_lengths]

        # 找出异常分片（超过告警阈值）
        abnormal_chunks = [(i, l) for i, l in chunk_lengths if l > WARNING_THRESHOLD]

        # 找出长度最大的N个分片
        top_largest = sorted(chunk_lengths, key=lambda x: x[1], reverse=True)[:MONITOR_CONFIG['log_top_largest_chunks']]

        return {
            'total': total_chunks,
            'max_length': max(lengths),
            'min_length': min(lengths),
            'avg_length': int(np.mean(lengths)),
            'abnormal_chunks': abnormal_chunks,
            'top_largest_chunks': top_largest
        }

    def get_text_embedding(self, text):
        """获取文本向量（优化批量处理，但保持接口兼容）"""
        try:
            if self.tokenizer is None or self.model is None:
                return [0.0] * 1024  # 默认向量

            inputs = self.tokenizer(text, return_tensors="pt", truncation=True, max_length=512)
            with torch.no_grad():
                outputs = self.model(**inputs)
                embedding = outputs.last_hidden_state.mean(dim=1).squeeze().numpy()
                return embedding.tolist()
        except Exception as e:
            logger.error(f"向量化失败: {str(e)}")
            return [0.0] * 1024

    def create_metadata_collection(self):
        """创建文档元数据集合"""
        try:
            if not connections.has_connection("default"):
                connections.connect(alias="default", host="localhost", port="19530")

            collection_name = "svc_bk_file_metadata"
            if collection_name in utility.list_collections():
                logger.info(f"集合已存在: {collection_name}")
                return True

            # 定义集合schema
            from pymilvus import FieldSchema, CollectionSchema

            fields = [
                FieldSchema(name="id", dtype=DataType.VARCHAR, is_primary=True, max_length=36),
                FieldSchema(name="bk_en_name", dtype=DataType.VARCHAR, max_length=100),
                FieldSchema(name="opportunity_number", dtype=DataType.VARCHAR, max_length=50),
                FieldSchema(name="original_filename", dtype=DataType.VARCHAR, max_length=255),
                FieldSchema(name="file_path", dtype=DataType.VARCHAR, max_length=500),
                FieldSchema(name="file_type", dtype=DataType.VARCHAR, max_length=10),
                FieldSchema(name="file_size", dtype=DataType.INT64),
                FieldSchema(name="upload_time", dtype=DataType.VARCHAR, max_length=20),
                FieldSchema(name="update_time", dtype=DataType.VARCHAR, max_length=20),
                FieldSchema(name="uploaded_by", dtype=DataType.VARCHAR, max_length=50),
                FieldSchema(name="kb_document_id", dtype=DataType.VARCHAR, max_length=36),
                FieldSchema(name="status", dtype=DataType.VARCHAR, max_length=20),
                FieldSchema(name="vector_dimension", dtype=DataType.INT64),
                FieldSchema(name="total_chunks", dtype=DataType.INT64),
                FieldSchema(name="chunk_strategy", dtype=DataType.VARCHAR, max_length=50),
                FieldSchema(name="error_msg", dtype=DataType.VARCHAR, max_length=500),
                FieldSchema(name="md5_hash", dtype=DataType.VARCHAR, max_length=32),
                FieldSchema(name="dummy_vector", dtype=DataType.FLOAT_VECTOR, dim=2),
            ]

            schema = CollectionSchema(fields, description="知识库文档元数据")
            collection = Collection(name=collection_name, schema=schema)

            # 为向量字段创建索引
            index_params = {
                "index_type": "IVF_FLAT",
                "metric_type": "L2",
                "params": {"nlist": 128}
            }
            collection.create_index("dummy_vector", index_params)

            logger.info(f"成功创建集合: {collection_name}")
            return True

        except Exception as e:
            logger.error(f"创建元数据集合失败: {str(e)}")
            return False

    def create_chunks_collection(self):
        """创建文档分片集合"""
        try:
            if not connections.has_connection("default"):
                connections.connect(alias="default", host="localhost", port="19530")

            collection_name = "svc_bk_file_vector_chunks"
            if collection_name in utility.list_collections():
                logger.info(f"集合已存在: {collection_name}")
                return True

            # 定义集合schema
            from pymilvus import FieldSchema, CollectionSchema

            fields = [
                FieldSchema(name="chunk_id", dtype=DataType.VARCHAR, is_primary=True, max_length=36),
                FieldSchema(name="document_id", dtype=DataType.VARCHAR, max_length=36),
                FieldSchema(name="chunk_index", dtype=DataType.INT64),
                FieldSchema(name="chunk_content", dtype=DataType.VARCHAR, max_length=MILVUS_CHUNK_CONTENT_MAX_LENGTH),
                FieldSchema(name="chunk_embedding", dtype=DataType.FLOAT_VECTOR, dim=1024),
                FieldSchema(name="dummy_vector", dtype=DataType.FLOAT_VECTOR, dim=2),
                FieldSchema(name="create_time", dtype=DataType.VARCHAR, max_length=20),
            ]

            schema = CollectionSchema(fields, description="知识库文档向量分片")
            collection = Collection(name=collection_name, schema=schema)

            # 为向量字段创建索引
            index_params = {
                "index_type": "IVF_FLAT",
                "metric_type": "L2",
                "params": {"nlist": 128}
            }
            collection.create_index("chunk_embedding", index_params)
            collection.create_index("dummy_vector", index_params)

            logger.info(f"成功创建集合: {collection_name}")
            return True

        except Exception as e:
            logger.error(f"创建分片集合失败: {str(e)}")
            return False

    def save_document_metadata(self, file_info):
        """保存文档元数据到svc_bk_file_metadata"""
        try:
            if not self.create_metadata_collection():
                return False

            # 连接到Milvus
            if not connections.has_connection("default"):
                connections.connect(alias="default", host="localhost", port="19530")

            collection = Collection("svc_bk_file_metadata")

            # 准备插入数据
            data = [
                [file_info["id"]],
                [file_info["bk_en_name"]],
                [file_info["opportunity_number"]],
                [file_info["original_filename"]],
                [file_info["file_path"]],
                [file_info["file_type"]],
                [file_info["file_size"]],
                [file_info["upload_time"]],
                [file_info["update_time"]],
                [file_info["uploaded_by"]],
                [file_info["kb_document_id"]],
                [file_info["status"]],
                [file_info["vector_dimension"]],
                [file_info["total_chunks"]],
                [file_info["chunk_strategy"]],
                [file_info["error_msg"]],
                [file_info["md5_hash"]],
                [[0.0, 0.0]],
            ]

            # 插入数据
            collection.insert(data)
            collection.flush()

            logger.info(f"成功保存文档元数据: {file_info['original_filename']}")
            return True

        except Exception as e:
            logger.error(f"保存文档元数据失败: {str(e)}")
            return False

    def save_document_chunks(self, document_id, chunks):
        """保存文档分片到svc_bk_file_vector_chunks（适配超大文件）"""
        try:
            if not self.create_chunks_collection():
                return {"success": False, "message": "创建分片集合失败"}

            # 连接到Milvus
            if not connections.has_connection("default"):
                connections.connect(alias="default", host="localhost", port="19530")

            collection = Collection("svc_bk_file_vector_chunks")

            # 分析分片长度（针对超大文件）
            chunk_analysis = self.analyze_chunk_lengths(chunks)
            logger.info(f"""
分片长度分析结果:
  - 总分片数: {chunk_analysis['total']}
  - 最大分片长度: {chunk_analysis['max_length']}
  - 平均分片长度: {chunk_analysis['avg_length']}
  - 异常分片数(>{WARNING_THRESHOLD}): {len(chunk_analysis['abnormal_chunks'])}
            """.strip())

            # 记录长度最大的分片信息
            if MONITOR_CONFIG['enable_extended_log'] and chunk_analysis['top_largest_chunks']:
                logger.info("长度最大的分片列表:")
                for idx, length in chunk_analysis['top_largest_chunks']:
                    logger.info(f"  分片 {idx}: {length} 字符")

            # 记录异常分片
            if MONITOR_CONFIG['log_abnormal_chunks'] and chunk_analysis['abnormal_chunks']:
                logger.warning("长度超过告警阈值的分片列表:")
                for idx, length in chunk_analysis['abnormal_chunks']:
                    logger.warning(f"  分片 {idx}: {length} 字符 (阈值: {WARNING_THRESHOLD})")

            # 最终解决方案：强制截断所有分片内容到安全长度
            safe_max_length = 1900  # 设置安全长度，留出100字符余量
            truncated_count = 0
            
            # 检查分片内容完整性
            logger.info("🔍 检查分片内容完整性...")
            
            # 准备批量插入数据
            batch_data = {
                'chunk_ids': [],
                'document_ids': [],
                'chunk_indexes': [],
                'chunk_contents': [],
                'chunk_embeddings': [],
                'dummy_vectors': [],
                'create_times': []
            }
            
            for i, chunk in enumerate(chunks):
                original_length = len(chunk["chunk_content"])
                
                # 检查第63个分片（错误信息中的问题分片）
                if i == 63:
                    logger.info(f"🔍 第63个分片原始长度: {original_length} 字符")
                    # 记录内容前100字符用于诊断
                    content_preview = chunk["chunk_content"][:100]
                    logger.info(f"   内容预览: {content_preview}")
                
                # 强制截断分片内容
                if original_length > safe_max_length:
                    chunk["chunk_content"] = chunk["chunk_content"][:safe_max_length]
                    truncated_count += 1
                    logger.warning(f"🔧 分片 {i} 强制截断: {original_length} -> {len(chunk['chunk_content'])} 字符")
                
                # 添加到批量数据
                batch_data['chunk_ids'].append(chunk["chunk_id"])
                batch_data['document_ids'].append(document_id)
                batch_data['chunk_indexes'].append(chunk["chunk_index"])
                batch_data['chunk_contents'].append(chunk["chunk_content"])
                batch_data['chunk_embeddings'].append(chunk["chunk_embedding"])
                batch_data['dummy_vectors'].append([0.0, 0.0])
                batch_data['create_times'].append(chunk["create_time"])
            
            # 记录截断统计
            if truncated_count > 0:
                logger.info(f"🔍 强制截断统计: {truncated_count}/{len(chunks)} 个分片被截断")
            else:
                logger.info("✅ 所有分片长度正常，无需截断")
            
            # 特别检查保存前的第63个分片内容
            if len(batch_data['chunk_contents']) > 63:
                final_length = len(batch_data['chunk_contents'][63])
                logger.info(f"🔍 保存前第63个分片最终长度: {final_length} 字符")
                if final_length > 2000:
                    logger.error(f"🚨 严重问题: 第63个分片保存前长度异常: {final_length} > 2000")
                    # 强制截断到安全长度
                    batch_data['chunk_contents'][63] = batch_data['chunk_contents'][63][:1900]
                    logger.warning(f"   已强制截断为: {len(batch_data['chunk_contents'][63])} 字符")

            # 输出截断统计
            logger.info(f"✅ 分片预处理完成: 共处理 {len(chunks)} 个分片，其中 {truncated_count} 个分片被截断")

            # 批量插入（针对超大文件，拆分批次插入）
            total_inserted = 0
            batch_size = 1000  # 每批次插入1000个分片，避免内存溢出
            total_batches = (len(batch_data['chunk_ids']) + batch_size - 1) // batch_size

            for batch_idx in range(total_batches):
                start_idx = batch_idx * batch_size
                end_idx = min((batch_idx + 1) * batch_size, len(batch_data['chunk_ids']))

                # 组装批次数据
                batch = [
                    batch_data['chunk_ids'][start_idx:end_idx],
                    batch_data['document_ids'][start_idx:end_idx],
                    batch_data['chunk_indexes'][start_idx:end_idx],
                    batch_data['chunk_contents'][start_idx:end_idx],
                    batch_data['chunk_embeddings'][start_idx:end_idx],
                    batch_data['dummy_vectors'][start_idx:end_idx],
                    batch_data['create_times'][start_idx:end_idx]
                ]

                # 插入批次数据
                collection.insert(batch)
                total_inserted += end_idx - start_idx

                # 输出批次进度
                logger.info(
                    f"📤 批次 {batch_idx + 1}/{total_batches} 插入完成: {end_idx - start_idx} 个分片 (累计: {total_inserted})")

            # 刷新数据
            collection.flush()

            logger.info(f"✅ 所有分片保存完成: 共插入 {total_inserted} 个分片到Milvus")
            return {
                "success": True,
                "message": f"成功保存 {total_inserted} 个分片",
                "truncated_count": truncated_count,
                "total_chunks": total_inserted
            }

        except Exception as e:
            error_msg = f"保存文档分片失败: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return {"success": False, "message": error_msg, "error_detail": str(e)}

    def process_document(self, file_path, kb_id, opportunity_number, original_filename, uploaded_by, document_id=None):
        """处理上传的文档（全面适配超大文件，支持Excel，自动清洗多媒体内容）"""
        try:
            logger.info(f"🚀 开始处理文档: {original_filename} (文件路径: {file_path})")

            # 记录文件大小（针对超大文件）
            file_size = os.path.getsize(file_path)
            file_size_mb = file_size / (1024 * 1024)
            logger.info(f"📁 文件大小: {file_size} 字节 ({file_size_mb:.2f} MB)")

            # 使用传入的document_id或生成新的
            if document_id:
                logger.info(f"✅ 使用传入的文档ID: {document_id}")
            else:
                document_id = str(uuid.uuid4())
                logger.info(f"✅ 生成新的文档ID: {document_id}")
            
            # 生成kb_document_id
            kb_document_id = str(uuid.uuid4())
            logger.info(f"✅ 生成kb_document_id: {kb_document_id}")

            # 获取文件信息
            logger.info("📊 步骤1: 获取文件信息...")
            file_ext = original_filename.lower().split('.')[-1]
            md5_hash = self.get_file_md5(file_path)
            logger.info(f"✅ 文件信息获取完成: 大小={file_size} bytes, 类型={file_ext}, MD5={md5_hash}")

            # 根据文件类型提取文本（支持PDF/Word/TXT/Excel）
            logger.info("📄 步骤2: 提取文本内容...")
            text = ""
            try:
                if file_ext == 'pdf':
                    text = self.extract_text_from_pdf(file_path)
                elif file_ext in ['doc', 'docx']:
                    text = self.extract_text_from_docx(file_path)
                elif file_ext in ['txt', 'sql']:
                    text = self.extract_text_from_txt(file_path)
                elif file_ext == 'xlsx':
                    # 检查openpyxl是否可用
                    if openpyxl is None:
                        error_msg = "openpyxl未安装，无法处理.xlsx文件，请执行: pip install openpyxl"
                        logger.error(error_msg)
                        return {"success": False, "step": "Excel处理", "message": error_msg}
                    text = self.extract_text_from_xlsx(file_path)
                elif file_ext == 'xls':
                    # 检查xlrd是否可用
                    if xlrd is None:
                        error_msg = "xlrd未安装，无法处理.xls文件，请执行: pip install xlrd==1.2.0"
                        logger.error(error_msg)
                        return {"success": False, "step": "Excel处理", "message": error_msg}
                    text = self.extract_text_from_xls(file_path)
                else:
                    error_msg = f"不支持的文件格式: {file_ext} (支持格式: pdf, doc, docx, txt, sql, xlsx, xls)"
                    logger.error(error_msg)
                    return {"success": False, "step": "文件格式检查", "message": error_msg}

                if not text:
                    error_msg = "文本提取失败或文件内容为空"
                    logger.error(error_msg)
                    return {"success": False, "step": "文本提取", "message": error_msg}
                    
            except Exception as e:
                error_msg = f"文件处理失败: {str(e)}"
                logger.error(error_msg)
                return {"success": False, "step": "文件处理", "message": error_msg}

            logger.info(f"✅ 文本提取完成，原始内容长度: {len(text)} 字符")

            # 核心新增：清洗图片/视频/多媒体内容
            logger.info("🧹 步骤2.5: 清洗图片/视频/Base64等多媒体内容...")
            clean_text = self.clean_media_content(text)
            logger.info(f"✅ 多媒体内容清洗完成，清洗后长度: {len(clean_text)} 字符")

            if not clean_text:
                error_msg = "清洗后文本内容为空（可能原文件仅包含图片/视频）"
                logger.error(error_msg)
                return {"success": False, "step": "多媒体清洗", "message": error_msg}

            # 文本分片（使用清洗后的纯文本）
            logger.info("✂️ 步骤3: 文本分片处理...")
            chunks = self.chunk_text(clean_text)
            if not chunks:
                error_msg = "文本分片失败"
                logger.error(error_msg)
                return {"success": False, "step": "文本分片", "message": error_msg}

            # 步骤3.5: 检查同名文档是否已存在，存在则删除旧数据（覆盖模式）
            logger.info("🔍 步骤3.5: 检查同名文档是否已存在...")
            overwrite = False
            try:
                if not connections.has_connection("default"):
                    connections.connect(alias="default", host="localhost", port="19530")
                if utility.has_collection("svc_bk_file_metadata"):
                    meta_col = Collection("svc_bk_file_metadata")
                    meta_col.load()
                    existing = meta_col.query(
                        expr=f'original_filename == "{original_filename}"',
                        output_fields=["id", "file_path"],
                        limit=100
                    )
                    if existing:
                        old_ids = [r["id"] for r in existing]
                        old_paths = [r.get("file_path", "") for r in existing]
                        logger.info(f"⚠️ 发现 {len(old_ids)} 条同名文档记录，执行覆盖删除...")
                        # 删除旧元数据
                        id_expr = ", ".join([f'"{i}"' for i in old_ids])
                        meta_col.delete(f'id in [{id_expr}]')
                        meta_col.flush()
                        # 删除旧向量分片
                        if utility.has_collection("svc_bk_file_vector_chunks"):
                            chunks_col = Collection("svc_bk_file_vector_chunks")
                            chunks_col.load()
                            chunks_col.delete(f'document_id in [{id_expr}]')
                            chunks_col.flush()
                            chunks_col.release()
                        meta_col.release()
                        # 删除旧磁盘文件
                        for old_path in old_paths:
                            if not old_path:
                                continue
                            if not os.path.isabs(old_path):
                                old_path = os.path.join(os.getcwd(), old_path)
                            if os.path.exists(old_path):
                                os.remove(old_path)
                                logger.info(f"🗑️ 已删除旧文件: {old_path}")
                            else:
                                logger.warning(f"⚠️ 旧文件不存在，跳过: {old_path}")
                        overwrite = True
                        logger.info("✅ 旧数据删除完成，准备写入新数据")
                    else:
                        meta_col.release()
            except Exception as e:
                logger.warning(f"⚠️ 查重检查异常（继续上传）: {str(e)}")

            # 准备文档元数据
            logger.info("📋 步骤4: 准备文档元数据...")
            now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            file_info = {
                "id": document_id,
                "bk_en_name": kb_id,
                "opportunity_number": opportunity_number,
                "original_filename": original_filename,
                "file_path": file_path,
                "file_type": file_ext,
                "file_size": file_size,
                "upload_time": now,
                "update_time": now,
                "uploaded_by": uploaded_by,
                "kb_document_id": kb_document_id,
                "status": "processing",
                "vector_dimension": 1024 if self.model else 0,
                "total_chunks": len(chunks),
                "chunk_strategy": f"char_500_ovlp_50_safe_{SAFE_TRUNCATE_LENGTH}_clean",
                "error_msg": "",
                "md5_hash": md5_hash
            }

            # 保存文档元数据
            logger.info("💾 步骤5: 保存文档元数据...")
            if not self.save_document_metadata(file_info):
                error_msg = "保存文档元数据失败"
                logger.error(error_msg)
                return {"success": False, "step": "保存元数据", "message": error_msg}

            # 处理文档分片（优化进度显示，适配大量分片）
            logger.info("🔍 步骤6: 处理文档分片...")
            processed_chunks = []
            total_chunks = len(chunks)
            progress_interval = max(1, total_chunks // 20)  # 最多显示20次进度

            for i, chunk in enumerate(chunks):
                chunk_id = str(uuid.uuid4())
                # 分片内容安全截断
                safe_chunk_content = self.safe_truncate_text(chunk)
                embedding = self.get_text_embedding(safe_chunk_content)

                chunk_info = {
                    "chunk_id": chunk_id,
                    "document_id": document_id,
                    "chunk_index": i,
                    "chunk_content": safe_chunk_content,
                    "chunk_embedding": embedding,
                    "create_time": now,
                }
                processed_chunks.append(chunk_info)

                # 智能进度显示（避免大量分片时日志刷屏）
                if (i + 1) % progress_interval == 0 or (i + 1) == total_chunks:
                    progress = (i + 1) / total_chunks * 100
                    logger.info(f"  进度: {i + 1}/{total_chunks} 分片处理完成 ({progress:.1f}%)")

            logger.info("✅ 文档分片处理完成")

            # 保存文档分片
            logger.info("💾 步骤7: 保存文档分片数据...")
            save_result = self.save_document_chunks(document_id, processed_chunks)
            logger.info(f"📊 save_document_chunks返回结果: {save_result}")

            if not save_result.get("success", False):
                error_msg = save_result.get("message", "保存文档分片失败")
                logger.error(f"🚨 保存文档分片失败: {error_msg}")
                return {"success": False, "step": "保存分片数据", "message": error_msg,
                        "error_detail": save_result.get("error_detail", "")}

            # 更新状态为完成
            file_info["status"] = "completed"
            file_info["update_time"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            self.save_document_metadata(file_info)

            logger.info(f"🎉 文档处理完成: {original_filename} (总分片: {len(processed_chunks)})")
            return {
                "success": True,
                "step": "完成",
                "message": f"文档处理成功（已清洗多媒体内容），共生成 {len(processed_chunks)} 个分片",
                "total_chunks": len(processed_chunks),
                "file_size_mb": file_size_mb,
                "cleaned_media": True,
                "overwrite": overwrite
            }

        except Exception as e:
            import traceback
            error_detail = traceback.format_exc()
            error_msg = f"文档处理失败: {str(e)}"
            logger.error(error_msg)
            logger.error(f"详细错误信息:\n{error_detail}")
            return {"success": False, "step": "未知", "message": error_msg, "error_detail": error_detail}


# 依赖安装提示
if __name__ == "__main__":
    print("""
    安装依赖（确保所有功能可用）：
    1. 基础依赖: pip install pymilvus PyPDF2 python-docx torch transformers numpy
    2. Excel处理: pip install openpyxl xlrd==1.2.0
    """)